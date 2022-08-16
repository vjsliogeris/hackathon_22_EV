from docx import Document
import os
import sys

from dotenv import dotenv_values

#Email imporrts

import smtplib
from os.path import basename
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.utils import COMMASPACE, formatdate
from email.mime.base import MIMEBase
import email.encoders as encoders

import translitcodec
import codecs

TEMPNAME = 'demptoc.docx'
SENDER_EMAIL = 'vytenis.sliogeris@matom.ai'

def handle_payslip(f) -> None:
    with open(TEMPNAME, 'wb+') as destination:
        for chunk in f.chunks():
            print(chunk)
            destination.write(chunk)
    # read sample docx and write to data.xml
    PATH = TEMPNAME

    print(TEMPNAME)
    document = Document(PATH)
    xml = document._body._body.xml
    xml_to_file = open("data.xml", "w")
    for line in xml:
        xml_to_file.writelines(line)
    xml_to_file.close()
    xml = open("data.xml", "r")
    ##

    read = False
    files = []
    temp = []
    while xml:
        line = xml.readline()
        line = line.rstrip()
        if line == "    <w:tr>": # start
            read = True
        if read:
            temp.append(line)
        if line == "    </w:tr>": # end
            read = False
            files.append(temp)
            temp = []
        if line == "":
            break
    xml.close()


    temp = {}
    docs = []
    next = False
    skip_next = False
    vardas = ""
    pavarde = ""
    next_key = None
    i = 0
    for file in files:
        for line in file:
            line = line.strip()
            if "<w:t>" in line:
              if 'EUR' in line:
                  continue
              elif 'UAB' in line:
                  continue
              elif 'ATSISKAITYMO' in line:
                  continue
              line = line[5:-6]
              if '· ' in line:
                  line = line[2:]
              if 'Laikotarpis' in line:
                  line = line[13:]
                  temp["Laikotarpis"] = line
                  next = True
              elif next:
                  if i == 0:
                      vardas = line
                      i += 1
                  elif i == 1:
                      pavarde = line
                      next = False
                      i = 0
                      temp["Darbuotojas"] = vardas + " " + pavarde
                      skip_next = True
              elif skip_next:
                  skip_next = False
                  continue
              elif next_key is None:
                  next_key = line
              elif next_key:
                  temp[next_key] = line
                  next_key = None

        docs.append(temp)
        temp = {}


    os.system("mkdir to_send")
    for doc in docs:
        i = 0
        output = Document()
        table = output.add_table(rows=len(doc), cols=2, style="Table Grid")
        for key in doc:
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = doc[key]
            i+=1
        to_send_fname = f"to_send/{doc['Darbuotojas']}.docx"
        print(to_send_fname)
        output.save(to_send_fname)
        darbotojas = doc['Darbuotojas']
        print(darbotojas)
        darbotojas = codecs.encode(darbotojas, 'translit/long')
        recipient = darbotojas.lower()
        recipient = recipient.replace(' ','.')
        recipient += "@matom.ai"
        print(recipient)
        send_mail(SENDER_EMAIL, recipient, to_send_fname)

    #cleanup
    os.system("rm data.xml")
    os.system("rm " + TEMPNAME)

    ### send

    #cleanup
    os.system("rm -rf to_send")


def send_mail(
        send_from: str,
        send_to: str,
        f: str,
        server: str = "smtp.gmail.com",
        port: int = 587) -> None:
    """Dedicated ,,spam direktoriai email'' funkcija

    mad props to the hood:
    https://stackoverflow.com/questions/3362600/how-to-send-email-attachments
    """
    denv = dotenv_values(".env")

    msg = MIMEMultipart()
    msg['From'] = send_from
    msg['To'] = send_to
    msg['Date'] = formatdate(localtime=True)
    msg['Subject'] = "You've got paid!"

    text = "Here comes the moolah"
    msg.attach(MIMEText(text, 'plain'))

    attachment = open(f, "rb")

    p = MIMEBase('application', 'octet-stream')
    p.set_payload(attachment.read())
    encoders.encode_base64(p)
    p.add_header('Content-Disposition', "attachment; filename= payslip.docx")
    msg.attach(p)

    app_pw = denv['APP_PW']
    s = smtplib.SMTP(server, port)
    s.starttls()
    s.login(send_from, app_pw)
    text = msg.as_string()
    s.sendmail(send_from, send_to, text)
    s.quit()
